import docx

doc = docx.Document('.\\chap 13\\demo.docx')
print(doc.paragraphs[1].text)
print(doc.paragraphs[0].text)